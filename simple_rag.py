"""
Concise RAG pipeline using LangChain + HuggingFace embeddings.
Reads 3 local Word docs, chunks & embeds them, stores vectors in FAISS (in-memory)
or Chroma (persistent), then runs a sample query using RetrievalQA.
"""

import os
from pathlib import Path
import docx

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS, Chroma
from langchain_community.llms import Ollama

# --- Config ---
FILES = ["doc1.docx", "doc2.docx", "doc3.docx"]  # local Word docs
VECTOR_STORE = "faiss"  # choose "faiss" or "chroma"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHROMA_PERSIST_DIR = "./chroma_store"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# --- 1. Load documents ---
docs = []
for file in FILES:
    text = "\n\n".join(p.text for p in docx.Document(file).paragraphs if p.text.strip())
    docs.append(Document(page_content=text))

# --- 2. Chunk documents ---
splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
chunks = []
for d in docs:
    for chunk in splitter.split_text(d.page_content):
        chunks.append(Document(page_content=chunk))
print(f"Total chunks: {len(chunks)}")

# --- 3. Embeddings ---
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

# --- 4. Vector store ---
if VECTOR_STORE == "faiss":
    print("Using in-memory FAISS store...")
    vs = FAISS.from_documents(chunks, embeddings)
elif VECTOR_STORE == "chroma":
    print(f"Using persistent Chroma store at {CHROMA_PERSIST_DIR} ...")
    os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
    vs = Chroma.from_documents(chunks, embeddings, persist_directory=CHROMA_PERSIST_DIR)
    vs.persist()
else:
    raise ValueError("VECTOR_STORE must be 'faiss' or 'chroma'")

# --- 5. Simple retrieval ---
llm = Ollama(model="mistral", temperature=0)  # Make sure Ollama is running: ollama run mistral
clearetriever = vs.as_retriever(search_kwargs={"k": 4})

# --- 6. Example query ---
query = "Summarize the main ideas from these documents."
print("\nQuery:", query)

# Retrieve relevant documents
retrieved_docs = retriever.invoke(query)
context = "\n\n".join([doc.page_content for doc in retrieved_docs])

# Create prompt and get answer
from langchain_core.prompts import ChatPromptTemplate
prompt = ChatPromptTemplate.from_template(
    "Based on the following context, answer the question.\n\n"
    "Context:\n{context}\n\n"
    "Question: {question}\n\n"
    "Answer:"
)

chain = prompt | llm
result = chain.invoke({"context": context, "question": query})
print("\nAnswer:\n", result.content)
